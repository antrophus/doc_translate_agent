import streamlit as st
from translator import DocumentTranslator
import os
from docx import Document
import tempfile
import logging
from styles import create_mobile_friendly_layout

# ë¡œê¹… ì„¤ì •
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def extract_text_from_table(table):
    """í‘œì—ì„œ í…ìŠ¤íŠ¸ë¥¼ ì¶”ì¶œí•©ë‹ˆë‹¤."""
    text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            text.append(" | ".join(row_text))
    return "\n".join(text)

def validate_docx(file_content):
    """ì—…ë¡œë“œëœ íŒŒì¼ì´ ìœ íš¨í•œ Word ë¬¸ì„œì¸ì§€ í™•ì¸í•©ë‹ˆë‹¤."""
    try:
        # ì„ì‹œ íŒŒì¼ë¡œ ì €ì¥
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name
        
        # ë¬¸ì„œ ì—´ê¸° ì‹œë„
        doc = Document(tmp_file_path)
        
        # ë¬¸ì„œ ë‚´ìš© ìƒì„¸ ë¡œê¹…
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        non_empty_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        
        logger.info(f"ë¬¸ì„œ ê²€ì¦ - ì „ì²´ ë‹¨ë½ ìˆ˜: {total_paragraphs}, í‘œ ìˆ˜: {total_tables}, ë‚´ìš©ì´ ìˆëŠ” ë‹¨ë½ ìˆ˜: {non_empty_paragraphs}")
        
        # í‘œ ë‚´ìš© ë¡œê¹…
        for i, table in enumerate(doc.tables):
            table_text = extract_text_from_table(table)
            if table_text:
                logger.info(f"í‘œ {i+1} ë‚´ìš© ìƒ˜í”Œ:\n{table_text[:500]}...")
        
        # ê° ë‹¨ë½ ë‚´ìš© ë¡œê¹…
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                logger.info(f"ë‹¨ë½ {i+1} ê¸¸ì´: {len(para.text)}, ë‚´ìš©: '{para.text[:100]}...'")
                for j, run in enumerate(para.runs):
                    logger.info(f"  - Run {j+1}: í…ìŠ¤íŠ¸: '{run.text}', ê¸€ê¼´: {run.font.name}")
        
        # ì„ì‹œ íŒŒì¼ ì‚­ì œ
        os.unlink(tmp_file_path)
        
        return total_tables > 0 or total_paragraphs > 0
    except Exception as e:
        logger.error(f"ë¬¸ì„œ ê²€ì¦ ì‹¤íŒ¨: {str(e)}")
        return False

def main():
    # ëª¨ë°”ì¼ ì¹œí™”ì ì¸ ë ˆì´ì•„ì›ƒ ìƒì„±
    container = create_mobile_friendly_layout()
    
    with container:
        st.title("ğŸ“ ë¬¸ì„œ ë²ˆì—­ê¸°")
        st.write("Word ë¬¸ì„œë¥¼ ì„ íƒí•œ ì–¸ì–´ë¡œ ë²ˆì—­í•©ë‹ˆë‹¤.")
        
        # ì ‘ì† ì •ë³´ í‘œì‹œ
        st.sidebar.info("""
        ### ì ‘ì† ì •ë³´
        - ë¡œì»¬ ì ‘ì†: http://localhost:8501
        - ë„¤íŠ¸ì›Œí¬ ì ‘ì†: http://192.168.0.80:8501
        - ëª¨ë°”ì¼ì—ì„œ ì ‘ì†í•˜ë ¤ë©´ ê°™ì€ WiFi ë„¤íŠ¸ì›Œí¬ì— ì—°ê²°ë˜ì–´ ìˆì–´ì•¼ í•©ë‹ˆë‹¤.
        """)
        
        # íŒŒì¼ ì—…ë¡œë“œ
        uploaded_file = st.file_uploader("ë²ˆì—­í•  Word ë¬¸ì„œë¥¼ ì—…ë¡œë“œí•˜ì„¸ìš”", type=['docx'])
        
        if uploaded_file is not None:
            # íŒŒì¼ í¬ê¸° í‘œì‹œ
            file_size = len(uploaded_file.getvalue())
            st.info(f"ì—…ë¡œë“œëœ íŒŒì¼ í¬ê¸°: {file_size / 1024:.2f} KB")
            
            # íŒŒì¼ ë‚´ìš© ë¯¸ë¦¬ë³´ê¸°
            try:
                # íŒŒì¼ í¬ì¸í„° ì´ˆê¸°í™”
                uploaded_file.seek(0)
                
                # ì„ì‹œ íŒŒì¼ë¡œ ì €ì¥
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                
                # ë¬¸ì„œ ì—´ê¸°
                doc = Document(tmp_file_path)
                
                # ë¬¸ì„œ ë‚´ìš© í™•ì¸
                paragraphs = []
                
                # í‘œ ë‚´ìš© ì¶”ì¶œ
                for table in doc.tables:
                    table_text = extract_text_from_table(table)
                    if table_text:
                        paragraphs.append(table_text)
                
                # ì¼ë°˜ ë‹¨ë½ ë‚´ìš© ì¶”ì¶œ
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                total_paragraphs = len(doc.paragraphs)
                total_tables = len(doc.tables)
                non_empty_content = len(paragraphs)
                
                logger.info(f"ë¬¸ì„œ ë¯¸ë¦¬ë³´ê¸° - ì „ì²´ ë‹¨ë½ ìˆ˜: {total_paragraphs}, í‘œ ìˆ˜: {total_tables}, ë‚´ìš©ì´ ìˆëŠ” ì„¹ì…˜ ìˆ˜: {non_empty_content}")
                
                if paragraphs:
                    preview_text = "\n\n".join(paragraphs[:5])
                    st.text_area("ë¬¸ì„œ ë¯¸ë¦¬ë³´ê¸° (ì²˜ìŒ 5ê°œ ì„¹ì…˜)", preview_text, height=300)
                    st.info(f"ì „ì²´ ë‹¨ë½ ìˆ˜: {total_paragraphs}, í‘œ ìˆ˜: {total_tables}, ë‚´ìš©ì´ ìˆëŠ” ì„¹ì…˜ ìˆ˜: {non_empty_content}")
                else:
                    st.warning("ë¬¸ì„œì— í…ìŠ¤íŠ¸ ë‚´ìš©ì´ ì—†ìŠµë‹ˆë‹¤. ë¬¸ì„œê°€ ì´ë¯¸ì§€ë‚˜ íŠ¹ìˆ˜ í˜•ì‹ìœ¼ë¡œë§Œ êµ¬ì„±ë˜ì–´ ìˆì„ ìˆ˜ ìˆìŠµë‹ˆë‹¤.")
                    # ë¬¸ì„œ êµ¬ì¡° ì •ë³´ í‘œì‹œ
                    st.info("ë¬¸ì„œ êµ¬ì¡° ì •ë³´:")
                    structure_info = []
                    for i, table in enumerate(doc.tables):
                        structure_info.append(f"í‘œ {i+1}: {len(table.rows)}í–‰ x {len(table.columns)}ì—´")
                    for i, para in enumerate(doc.paragraphs):
                        structure_info.append(f"ë‹¨ë½ {i+1}: ê¸¸ì´={len(para.text)}, ìŠ¤íƒ€ì¼={para.style.name}")
                    st.code("\n".join(structure_info))
                
                # ì„ì‹œ íŒŒì¼ ì‚­ì œ
                os.unlink(tmp_file_path)
                
            except Exception as e:
                logger.error(f"ë¬¸ì„œ ë¯¸ë¦¬ë³´ê¸° ì‹¤íŒ¨: {str(e)}")
                st.error(f"ë¬¸ì„œ ë¯¸ë¦¬ë³´ê¸° ì‹¤íŒ¨: {str(e)}")
        
        # ë²ˆì—­í•  ì–¸ì–´ ì„ íƒ
        target_lang = st.selectbox(
            "ë²ˆì—­í•  ì–¸ì–´ë¥¼ ì„ íƒí•˜ì„¸ìš”",
            ["ì¤‘êµ­ì–´ ê°„ì²´", "ì˜ì–´", "ì¼ë³¸ì–´", "ë² íŠ¸ë‚¨ì–´", "íƒœêµ­ì–´", "ì¸ë„ë„¤ì‹œì•„ì–´"]
        )
        
        if uploaded_file is not None and st.button("ë²ˆì—­ ì‹œì‘", type="primary"):
            try:
                # íŒŒì¼ ê²€ì¦
                if not validate_docx(uploaded_file.getvalue()):
                    st.error("ìœ íš¨í•˜ì§€ ì•Šì€ Word ë¬¸ì„œì…ë‹ˆë‹¤. ë‚´ìš©ì´ ìˆëŠ” Word ë¬¸ì„œë¥¼ ì—…ë¡œë“œí•´ì£¼ì„¸ìš”.")
                    return
                
                # ì„ì‹œ íŒŒì¼ë¡œ ì—…ë¡œë“œëœ íŒŒì¼ ì €ì¥
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                
                # ë²ˆì—­ê¸° ì´ˆê¸°í™”
                translator = DocumentTranslator()
                
                # ì§„í–‰ ìƒíƒœ í‘œì‹œ
                with st.spinner("ë²ˆì—­ ì¤‘ì…ë‹ˆë‹¤..."):
                    # ë¬¸ì„œ ë²ˆì—­
                    translated_doc = translator.translate_document(tmp_file_path, target_lang)
                    
                    # ë²ˆì—­ëœ ë¬¸ì„œ ì €ì¥
                    output_filename = f"translated_{uploaded_file.name}"
                    output_path = os.path.join("data", output_filename)
                    
                    # data ë””ë ‰í† ë¦¬ê°€ ì—†ìœ¼ë©´ ìƒì„±
                    os.makedirs("data", exist_ok=True)
                    
                    # ë²ˆì—­ëœ ë¬¸ì„œ ì €ì¥
                    translated_doc.save(output_path)
                    st.success(f"ë²ˆì—­ì´ ì™„ë£Œë˜ì—ˆìŠµë‹ˆë‹¤. íŒŒì¼ì´ ë‹¤ìŒ ê²½ë¡œì— ì €ì¥ë˜ì—ˆìŠµë‹ˆë‹¤: {output_path}")
                    
                    # ë²ˆì—­ëœ íŒŒì¼ ë‹¤ìš´ë¡œë“œ ë²„íŠ¼
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label="ë²ˆì—­ëœ ë¬¸ì„œ ë‹¤ìš´ë¡œë“œ",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    # ì„ì‹œ íŒŒì¼ ì‚­ì œ
                    os.unlink(tmp_file_path)
                    
            except Exception as e:
                logger.error(f"ë²ˆì—­ ì¤‘ ì˜¤ë¥˜ ë°œìƒ: {str(e)}")
                st.error(f"ë²ˆì—­ ì¤‘ ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤: {str(e)}")

if __name__ == "__main__":
    main() 