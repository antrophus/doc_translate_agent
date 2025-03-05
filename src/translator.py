from typing import Optional, Dict, Any
from langchain_community.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.callbacks.manager import get_openai_callback
import google.generativeai as genai
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
import streamlit as st
import logging
import time
from datetime import datetime, timedelta

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 환경 변수 로드
load_dotenv()

class DocumentTranslator:
    def __init__(self):
        logger.info("DocumentTranslator 초기화 시작")
        # API 키 설정
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
        
        logger.info("API 키 로드 완료")
        
        # OpenAI 설정 (최신 공식 문서 기준)
        self.openai_client = OpenAI()  # API 키는 환경 변수에서 자동으로 로드
        
        # Gemini 설정 (최신 공식 문서 기준)
        genai.configure(api_key=self.gemini_api_key)
        self.gemini_model = genai.GenerativeModel('gemini-2.0-flash')
        
        # DeepSeek 설정
        self.deepseek_client = OpenAI(
            api_key=self.deepseek_api_key,
            base_url="https://api.deepseek.com/v1"
        )
        
        logger.info("API 클라이언트 초기화 완료")
        
        # 번역 시스템 프롬프트
        self.system_prompt = """당신은 전문 번역가입니다. 
        다음 규칙을 엄격히 지켜주세요:
        1. 주어진 한국어 텍스트를 정확하고 자연스럽게 번역해주세요.
        2. 전문 용어나 문맥을 고려하여 번역해주세요.
        3. 번역할 때는 원문의 의미와 뉘앙스를 최대한 살려주세요.
        4. 번역 결과만 출력해주세요. 다른 설명이나 주석을 추가하지 마세요.
        5. 번역이 불가능한 경우 "번역 불가"라고만 답변해주세요.
        6. 번역할 때는 원문의 형식(줄바꿈, 들여쓰기, 표 형식 등)을 유지해주세요."""
    
    def translate_with_openai(self, text: str, target_lang: str) -> str:
        """OpenAI를 사용하여 텍스트를 번역합니다."""
        try:
            logger.info(f"OpenAI API 호출 시작 - 텍스트 길이: {len(text)}")
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": f"다음 한국어 텍스트를 {target_lang}로 번역해주세요:\n\n{text}"}
                ]
            )
            result = response.choices[0].message.content.strip()
            logger.info(f"OpenAI API 응답 완료 - 번역 결과 길이: {len(result)}")
            return result
        except Exception as e:
            logger.error(f"OpenAI API 호출 실패: {str(e)}")
            st.error(f"OpenAI 번역 중 오류 발생: {str(e)}")
            return None
    
    def translate_with_gemini(self, text: str, target_lang: str) -> str:
        """Gemini를 사용하여 텍스트를 번역합니다."""
        try:
            logger.info(f"Gemini API 호출 시작 - 텍스트 길이: {len(text)}")
            response = self.gemini_model.generate_content(
                f"{self.system_prompt}\n\n다음 한국어 텍스트를 {target_lang}로 번역해주세요:\n\n{text}"
            )
            result = response.text.strip()
            logger.info(f"Gemini API 응답 완료 - 번역 결과 길이: {len(result)}")
            return result
        except Exception as e:
            logger.error(f"Gemini API 호출 실패: {str(e)}")
            st.error(f"Gemini 번역 중 오류 발생: {str(e)}")
            return None
    
    def translate_with_deepseek(self, text: str, target_lang: str) -> str:
        """DeepSeek를 사용하여 텍스트를 번역합니다."""
        try:
            logger.info(f"DeepSeek API 호출 시작 - 텍스트 길이: {len(text)}")
            response = self.deepseek_client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": f"다음 한국어 텍스트를 {target_lang}로 번역해주세요:\n\n{text}"}
                ],
                temperature=0.7
            )
            result = response.choices[0].message.content.strip()
            logger.info(f"DeepSeek API 응답 완료 - 번역 결과 길이: {len(result)}")
            return result
        except Exception as e:
            logger.error(f"DeepSeek API 호출 실패: {str(e)}")
            st.error(f"DeepSeek 번역 중 오류 발생: {str(e)}")
            return None
    
    def translate_text(self, text: str, target_lang: str) -> str:
        """Failover 메커니즘을 사용하여 텍스트를 번역합니다."""
        if not text.strip():
            logger.warning("빈 텍스트 입력됨")
            return ""
            
        logger.info(f"번역 시작 - 대상 언어: {target_lang}")
            
        # 중국어 번역의 경우 DeepSeek를 우선 사용
        if target_lang == "중국어 간체":
            logger.info("DeepSeek API로 중국어 번역 시도")
            result = self.translate_with_deepseek(text, target_lang)
            if result and result != "번역 불가":
                logger.info("DeepSeek API 번역 성공")
                return result
        
        # 다른 언어의 경우 OpenAI를 우선 사용
        logger.info("OpenAI API로 번역 시도")
        result = self.translate_with_openai(text, target_lang)
        if result and result != "번역 불가":
            logger.info("OpenAI API 번역 성공")
            return result
        
        # OpenAI 실패시 Gemini로 시도
        logger.info("Gemini API로 번역 시도")
        result = self.translate_with_gemini(text, target_lang)
        if result and result != "번역 불가":
            logger.info("Gemini API 번역 성공")
            return result
        
        # Gemini 실패시 DeepSeek로 시도
        logger.info("DeepSeek API로 마지막 번역 시도")
        result = self.translate_with_deepseek(text, target_lang)
        if result and result != "번역 불가":
            logger.info("DeepSeek API 번역 성공")
            return result
        
        logger.error("모든 번역 서비스 실패")
        raise Exception("모든 번역 서비스가 실패했습니다.")
    
    def translate_document(self, doc_path: str, target_lang: str) -> Document:
        """Word 문서를 번역합니다."""
        start_time = time.time()
        logger.info(f"문서 번역 시작 - 파일: {doc_path}, 대상 언어: {target_lang}")
        doc = Document(doc_path)
        
        # 진행 상태 표시를 위한 프로그레스 바와 상태 텍스트
        progress_bar = st.progress(0)
        status_text = st.empty()
        time_info = st.empty()
        completion_info = st.empty()
        
        def update_time_info(progress: float, start_time: float, is_completed: bool = False):
            """진행 시간과 예상 완료 시간을 업데이트합니다."""
            elapsed_time = time.time() - start_time
            if is_completed:
                time_info.text(f"총 소요 시간: {timedelta(seconds=int(elapsed_time))}")
                completion_info.success("번역이 완료되었습니다!")
            elif progress > 0:
                estimated_total_time = elapsed_time / progress
                remaining_time = estimated_total_time - elapsed_time
                estimated_completion = datetime.now() + timedelta(seconds=remaining_time)
                
                time_info.text(
                    f"진행 시간: {timedelta(seconds=int(elapsed_time))} | "
                    f"예상 남은 시간: {timedelta(seconds=int(remaining_time))} | "
                    f"예상 완료 시간: {estimated_completion.strftime('%H:%M:%S')}"
                )
        
        def is_translatable(text: str) -> bool:
            """번역이 필요한 텍스트인지 확인합니다."""
            # 숫자나 특수문자만 있는 경우 번역하지 않음
            if text.strip().replace('.', '').replace(',', '').replace(' ', '').isdigit():
                return False
            # 한글이나 의미있는 문자가 포함된 경우에만 번역
            return any(ord('가') <= ord(c) <= ord('힣') for c in text)
        
        def batch_translate_cells(cells_data: list) -> list:
            """여러 셀의 텍스트를 한 번에 번역합니다."""
            if not cells_data:
                return []
                
            try:
                # 셀 텍스트를 하나의 문자열로 결합 (구분자로 구분)
                combined_text = "\n---CELL_SEPARATOR---\n".join([text for _, text in cells_data])
                translated_text = self.translate_text(combined_text, target_lang)
                
                if not translated_text:
                    logger.warning("번역된 텍스트가 비어있습니다.")
                    return [""] * len(cells_data)
                    
                # 번역된 텍스트를 다시 개별 셀로 분리
                translated_parts = translated_text.split("---CELL_SEPARATOR---")
                
                # 원본 셀 수와 번역된 부분의 수가 일치하지 않을 경우 처리
                if len(translated_parts) != len(cells_data):
                    logger.warning(f"번역된 셀 수가 일치하지 않습니다. 원본: {len(cells_data)}, 번역: {len(translated_parts)}")
                    # 부족한 부분은 빈 문자열로 채움
                    if len(translated_parts) < len(cells_data):
                        translated_parts.extend([""] * (len(cells_data) - len(translated_parts)))
                    # 초과된 부분은 제거
                    translated_parts = translated_parts[:len(cells_data)]
                    
                return [part.strip() for part in translated_parts]
                
            except Exception as e:
                logger.error(f"일괄 번역 중 오류 발생: {str(e)}")
                return [""] * len(cells_data)
        
        # 표 내용 번역
        for table_idx, table in enumerate(doc.tables):
            logger.info(f"표 {table_idx + 1} 번역 시작")
            translatable_cells = []
            
            # 번역이 필요한 셀 수집
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and is_translatable(cell.text):
                        translatable_cells.append((cell, cell.text))
            
            # 배치 크기 설정 (한 번에 10개 셀씩 처리)
            BATCH_SIZE = 10
            total_batches = (len(translatable_cells) + BATCH_SIZE - 1) // BATCH_SIZE
            
            for batch_idx in range(total_batches):
                start_idx = batch_idx * BATCH_SIZE
                end_idx = min((batch_idx + 1) * BATCH_SIZE, len(translatable_cells))
                current_batch = translatable_cells[start_idx:end_idx]
                
                # 배치 번역 실행
                translated_texts = batch_translate_cells(current_batch)
                
                # 번역 결과 적용
                for i, translated_text in enumerate(translated_texts):
                    if translated_text and translated_text.strip():
                        current_batch[i][0].text = translated_text.strip()
                
                # 진행 상태 업데이트
                progress = (table_idx + (batch_idx + 1) / total_batches) / (len(doc.tables) + 1)  # +1은 단락 처리를 위한 여유
                progress_bar.progress(progress)
                status_text.text(f"표 {table_idx + 1}/{len(doc.tables)} 번역 중... ({batch_idx + 1}/{total_batches} 배치)")
                update_time_info(progress, start_time)
        
        # 일반 단락 번역
        paragraphs_to_translate = [p for p in doc.paragraphs if p.text.strip() and is_translatable(p.text)]
        total_paragraphs = len(paragraphs_to_translate)
        
        for i, paragraph in enumerate(paragraphs_to_translate):
            if paragraph.text.strip():
                try:
                    translated_text = self.translate_text(paragraph.text, target_lang)
                    if translated_text and translated_text != "번역 불가":
                        paragraph.text = translated_text
                    progress = (len(doc.tables) + (i + 1) / total_paragraphs) / (len(doc.tables) + 1)
                    progress_bar.progress(progress)
                    status_text.text(f"단락 번역 중... ({i + 1}/{total_paragraphs})")
                    update_time_info(progress, start_time)
                except Exception as e:
                    logger.error(f"단락 번역 중 오류 발생: {str(e)}")
                    continue
        
        logger.info("문서 번역 완료")
        # 번역 완료 시 상태 업데이트
        progress_bar.progress(1.0)
        status_text.text("번역 완료")
        update_time_info(1.0, start_time, is_completed=True)
        return doc 